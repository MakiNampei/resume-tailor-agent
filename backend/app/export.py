import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from typing import Dict, Any, List

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# --- Helpers ---

def get_ordered_categories() -> List[str]:
    return ["education", "work", "project", "leadership"]

def get_category_display_name(category: str) -> str:
    mapping = {
        "education": "EDUCATION",
        "work": "WORK EXPERIENCE",
        "project": "PROJECTS & EXPERIENCES",
        "leadership": "LEADERSHIP & ACTIVITIES"
    }
    return mapping.get(category.lower(), category.upper())

def group_blocks_by_category(selected_blocks: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    grouped = {}
    for block in selected_blocks:
        cat = block.get("category", "work").lower()
        if cat not in grouped:
            grouped[cat] = []
            
        # Extract tailored bullets or fallback to original
        bullets = block.get("tailored_bullets")
        if not bullets:
            bullets = block.get("original_bullets", [])
            
        grouped[cat].append({
            "title": block.get("title", ""),
            "organization": block.get("organization", ""),
            "location": block.get("location", ""),
            "start_date": block.get("start_date", ""),
            "end_date": block.get("end_date", ""),
            "bullets": bullets
        })
    return grouped


# --- DOCX Generation ---

def generate_docx_resume(resume_data: Dict[str, Any]) -> io.BytesIO:
    doc = Document()
    
    # 1. Page Margins (0.75 in all around is standard and professional)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Set default font to Calibri or Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    
    personal_info = resume_data.get("personal_info", {})
    selected_blocks = resume_data.get("selected_blocks", [])
    
    # 2. Header
    header_p = doc.add_paragraph()
    header_p.alignment = 1  # Centered
    header_p.paragraph_format.space_after = Pt(2)
    
    name_run = header_p.add_run(personal_info.get("full_name", "Your Name") + "\n")
    name_run.font.size = Pt(18)
    name_run.bold = True
    
    contact_parts = []
    if personal_info.get("email"):
        contact_parts.append(personal_info["email"])
    if personal_info.get("phone"):
        contact_parts.append(personal_info["phone"])
    if personal_info.get("location"):
        contact_parts.append(personal_info["location"])
        
    links = []
    if personal_info.get("linkedin"):
        links.append("LinkedIn: " + personal_info["linkedin"].replace("https://", "").replace("www.", ""))
    if personal_info.get("github"):
        links.append("GitHub: " + personal_info["github"].replace("https://", "").replace("www.", ""))
    if personal_info.get("website"):
        links.append(personal_info["website"].replace("https://", "").replace("www.", ""))
        
    contact_text = "  |  ".join(contact_parts)
    if links:
        contact_text += "\n" + "  |  ".join(links)
        
    contact_run = header_p.add_run(contact_text)
    contact_run.font.size = Pt(9.5)
    
    # Add a thin line spacing after header
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # 3. Resume Sections
    grouped_blocks = group_blocks_by_category(selected_blocks)
    
    # Use tab stop at 7.0 inches (since printable area is 8.5" width - 2 * 0.75" margins = 7.0" printable width)
    tab_stop_pos = Inches(7.0)
    
    for category in get_ordered_categories():
        if category not in grouped_blocks or not grouped_blocks[category]:
            continue
            
        # Section Heading
        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(12)
        heading_p.paragraph_format.space_after = Pt(3)
        heading_p.paragraph_format.keep_with_next = True
        
        heading_run = heading_p.add_run(get_category_display_name(category))
        heading_run.font.size = Pt(11)
        heading_run.bold = True
        
        # Add a subtle bottom border or underline using a solid line
        # A clean way in DOCX is to put a bottom border on the paragraph, but a simple text divider is more compatible
        # Let's keep it clean: just a line of characters or a paragraph line
        divider_p = doc.add_paragraph()
        divider_p.paragraph_format.space_before = Pt(0)
        divider_p.paragraph_format.space_after = Pt(4)
        divider_p.paragraph_format.keep_with_next = True
        div_run = divider_p.add_run("―" * 68)  # standard width line
        div_run.font.size = Pt(6)
        div_run.font.color.rgb = docx.shared.RGBColor(160, 160, 160) if 'docx' in globals() else None
        
        for block in grouped_blocks[category]:
            block_p = doc.add_paragraph()
            block_p.paragraph_format.space_before = Pt(4)
            block_p.paragraph_format.space_after = Pt(2)
            block_p.paragraph_format.tab_stops.add_tab_stop(tab_stop_pos, WD_TAB_ALIGNMENT.RIGHT)
            block_p.paragraph_format.keep_with_next = True
            
            # Left side: Bold Title, Italic Org, Location
            title_run = block_p.add_run(block["title"])
            title_run.bold = True
            
            block_p.add_run("  |  ")
            
            org_run = block_p.add_run(block["organization"])
            org_run.italic = True
            
            if block.get("location"):
                block_p.add_run(f", {block['location']}")
                
            # Right side: Dates (Separated by Tab)
            date_str = f"\t{block['start_date']} – {block['end_date']}"
            date_run = block_p.add_run(date_str)
            date_run.font.size = Pt(10)
            
            # Bullets
            for bullet in block["bullets"]:
                bullet_p = doc.add_paragraph(style='List Bullet')
                bullet_p.paragraph_format.space_before = Pt(0)
                bullet_p.paragraph_format.space_after = Pt(2)
                
                b_run = bullet_p.add_run(bullet)
                b_run.font.name = 'Calibri'
                b_run.font.size = Pt(10)
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# --- PDF Generation ---

def generate_pdf_resume(resume_data: Dict[str, Any]) -> io.BytesIO:
    file_stream = io.BytesIO()
    
    # Letter size page is 8.5 x 11 inches. With 0.75" margins, width is 7.0 inches.
    doc = SimpleDocTemplate(
        file_stream,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Styles (using Helvetica for standard clean ATS look)
    name_style = ParagraphStyle(
        'ResumeName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        alignment=1, # Center
        spaceAfter=2
    )
    
    contact_style = ParagraphStyle(
        'ResumeContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        alignment=1, # Center
        spaceAfter=12
    )
    
    section_heading_style = ParagraphStyle(
        'ResumeSectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#1f2937'),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    title_style = ParagraphStyle(
        'ResumeBlockTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=13
    )
    
    org_style = ParagraphStyle(
        'ResumeBlockOrg',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=10,
        leading=13
    )
    
    date_style = ParagraphStyle(
        'ResumeBlockDate',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        alignment=2 # Right
    )
    
    bullet_style = ParagraphStyle(
        'ResumeBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=13,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=2
    )
    
    story = []
    
    personal_info = resume_data.get("personal_info", {})
    selected_blocks = resume_data.get("selected_blocks", [])
    
    # 1. Header Name
    story.append(Paragraph(personal_info.get("full_name", "Your Name"), name_style))
    
    # 2. Contact Info
    contact_parts = []
    if personal_info.get("email"):
        contact_parts.append(personal_info["email"])
    if personal_info.get("phone"):
        contact_parts.append(personal_info["phone"])
    if personal_info.get("location"):
        contact_parts.append(personal_info["location"])
        
    links = []
    if personal_info.get("linkedin"):
        links.append("LinkedIn: " + personal_info["linkedin"].replace("https://", "").replace("www.", ""))
    if personal_info.get("github"):
        links.append("GitHub: " + personal_info["github"].replace("https://", "").replace("www.", ""))
    if personal_info.get("website"):
        links.append(personal_info["website"].replace("https://", "").replace("www.", ""))
        
    contact_text = "  |  ".join(contact_parts)
    if links:
        contact_text += "<br/>" + "  |  ".join(links)
        
    story.append(Paragraph(contact_text, contact_style))
    
    # 3. Sections
    grouped_blocks = group_blocks_by_category(selected_blocks)
    
    for category in get_ordered_categories():
        if category not in grouped_blocks or not grouped_blocks[category]:
            continue
            
        # Draw Section Header with horizontal line below it using a 1-column Table
        heading_text = get_category_display_name(category)
        heading_table = Table(
            [[Paragraph(heading_text, section_heading_style)]],
            colWidths=[7.0*inch]
        )
        heading_table.setStyle(TableStyle([
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('LINEBELOW', (0,0), (-1,-1), 0.75, colors.HexColor('#9ca3af')),
        ]))
        
        story.append(heading_table)
        story.append(Spacer(1, 4))
        
        for block in grouped_blocks[category]:
            # Construct Block Header Table: Left has Title | Org, Right has Dates
            left_text = f"<b>{block['title']}</b> &nbsp;|&nbsp; <i>{block['organization']}</i>"
            if block.get("location"):
                left_text += f", {block['location']}"
                
            left_para = Paragraph(left_text, title_style)
            right_para = Paragraph(f"{block['start_date']} &ndash; {block['end_date']}", date_style)
            
            # Left gets 5.3 inches, right gets 1.7 inches
            header_table = Table(
                [[left_para, right_para]],
                colWidths=[5.3*inch, 1.7*inch]
            )
            header_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 3),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            
            story.append(header_table)
            
            # Bullets
            for bullet in block["bullets"]:
                story.append(Paragraph(f"&bull; {bullet}", bullet_style))
                
            story.append(Spacer(1, 4))
            
    doc.build(story)
    file_stream.seek(0)
    return file_stream
